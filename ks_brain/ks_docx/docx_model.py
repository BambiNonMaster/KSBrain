import os.path
import pandas as pd
from .docx_setting import WORD_SETTING, ALIGNMENT_DICT, HEADING_NUMBER, TEXT
from docx import Document
from docx.shared import (
    Pt,                                 # 磅值
    Cm,                                 # 厘米单位
    Length,
    RGBColor
)
from docx.oxml.ns import qn                                 # 中文字体
import re
from os import remove, listdir


"""
author: mengru.du
create: 2022.3.15
modify: 20220315

该文件定义文档输出基本框架类
"""


class DocxModel():
    def __init__(self, config: dict):
        """ 创建文件 """
        self.new_docx = Document()
        self.config = config

    def write_document(self, save_file: str):
        # 1. 设置正文格式
        self.new_docx.styles["Normal"].font.name = WORD_SETTING["英文字体"]
        self.new_docx.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_SETTING["中文字体"])
        self.new_docx.styles["Normal"].font.size = Pt(WORD_SETTING["正文字号"])
        self.new_docx.styles["Normal"].font.bold = WORD_SETTING["正文加粗"]

        # 2. 添加内容
        self.new_docx.add_heading(text=TEXT["head_1"], level=1)
        self.new_docx.add_heading(text=TEXT["head_1.1"], level=2)
        self.new_docx.add_heading(text=TEXT["head_1.1.1"], level=3)
        for para in TEXT["para_1.1.1"].split("\n"):
            self.new_docx.add_paragraph(text=para)
        self.new_docx.add_heading(text=TEXT["head_1.1.2"], level=3)
        for para in TEXT["para_1.1.2"].split("\n"):
            self.new_docx.add_paragraph(text=para)

        # ----- 第二章: 因子筛选汇总 ----- #
        self.new_docx.add_heading(text=TEXT["head_1.2"], level=2)
        # 添加表格(内容被封装了)
        self.add_table_factors()

        # ----- 第三章: 模型结果汇总 ----- #
        self.new_docx.add_heading(text=TEXT["head_1.3"], level=2)
        # 添加表格(内容被封装了)
        self.add_talbe_results()

        # ----- 第三章: 模型结果展示 ----- #
        self.new_docx.add_heading(text=TEXT["head_1.4"], level=2)
        # 添加图片
        self.add_pictures()

        self.set_format()
        # 3. 保存文件
        self.new_docx.save(save_file)

    def add_table_factors(self):
        """
         添加表格, 表格样式太复杂, 还是建议手动微调
         @rtype: object
         """
        factors = {}
        for models_env in self.config["models_env"]:
            active = models_env["active"]
            if not active:
                continue
            file = models_env["factors_file"]
            df_factors = pd.read_csv(f"f_output/{file}.csv", index_col=0)
            factors["模型"] = file
            factors["因子"] = [", ".join(df_factors.columns.to_list()[1:])]

        df_table = pd.DataFrame(factors)
        rows, cols = df_table.shape
        rows += 1  # 因为要添加标题, 所以行+1
        table = self.new_docx.add_table(rows=rows, cols=cols, style=WORD_SETTING["表格样式"])
        for row in range(rows):
            for col in range(cols):
                cell = table.cell(row, col)
                if row == 0:  # 填充标题
                    cell.text = str(df_table.columns.to_list()[col])
                else:  # 因为row第一行是标题, 所以row取值要-1
                    cell.text = str(df_table.iloc[row - 1, col])
                # 单独设置格式
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)  # 字体大小设置，和word里面的字号相对应
                        # run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色设置，这里是用RGB颜色
        # 调整行高
        for row in range(rows):
            table.rows[row].height = Cm(0.8)

        # 左对齐
        table.style.paragraph_format.alignment = ALIGNMENT_DICT[-1]

    def add_talbe_results(self):
        """
        添加表格, 表格样式太复杂, 还是建议手动微调
        @rtype: object
        """
        df_table = pd.read_csv(f"m_output/{self.config['symbol']}_accuracy.csv")
        rows, cols = df_table.shape
        rows += 1   # 因为要添加标题, 所以行+1
        table = self.new_docx.add_table(rows=rows, cols=cols, style=WORD_SETTING["表格样式"])
        for row in range(rows):
            for col in range(cols):
                cell = table.cell(row, col)
                if row == 0:            # 填充标题
                    cell.text = str(df_table.columns.to_list()[col])
                else:                   # 因为row第一行是标题, 所以row取值要-1
                    cell.text = str(df_table.iloc[row-1, col])
                # 单独设置格式
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)                      # 字体大小设置，和word里面的字号相对应
                        # run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色设置，这里是用RGB颜色
        # 调整行高
        for row in range(rows):
            table.rows[row].height = Cm(0.8)
        # 合并表格
        # cell_1 = table.cell(1, 0)
        # cell_2 = table.cell(2, 0)
        # cell_3 = table.cell(3, 0)
        # cell_4 = table.cell(4, 0)
        # cell_1.merge(cell_2).merge(cell_3).merge(cell_4)

        # 居中对齐
        table.style.paragraph_format.alignment = ALIGNMENT_DICT[0]

    def add_pictures(self):
        """
        添加图片
        @rtype: object
        """
        algo_count = 1      # 记录算法数量
        for env in self.config["models_env"]:
            if env["active"]:
                factors_file = env["factors_file"]
                # 1. 添加算法标题
                algo_name = factors_file.split("_")[0]
                self.new_docx.add_heading(text=f"1.3.{algo_count} 算法-{algo_name}", level=3)

                model_count = 1  # 记录模型数量
                for models in env["models"]:
                    add_model = models["add_model"]
                    model_name = add_model["model_name"]
                    # 2. 添加模型标题
                    self.new_docx.add_heading(text=f"1.3.{algo_count}.{model_count} {model_name}", level=4)

                    # 3. 添加模型正文
                    train_params = add_model["train_params"]
                    self.new_docx.add_paragraph(
                        f"模型参数: \n"
                        f"预测期数: {train_params['cadence']}, 训练频率: {train_params['retrain_freq']}, 数据量: {train_params['train_window']}")

                    # 4. 添加模型图片
                    new_para = self.new_docx.add_paragraph()
                    new_para.add_run().add_picture(f"m_images/{factors_file}_{model_name}.png",
                                                   width=Cm(WORD_SETTING["图片宽度"]), height=Cm(WORD_SETTING["图片高度"]))
                    model_count += 1
                algo_count += 1

    def set_format(self):
        """
        设置文本格式
        """
        # 最后调整段落格式
        for para in self.new_docx.paragraphs:
            if para.style.name == "Normal":
                para.paragraph_format.space_before = 1
                para.paragraph_format.space_after = 1
                # para.paragraph_format.line_spacing = 0            # 段间
                # para.paragraph_format.first_line_indent = Pt(0)   # 首行缩进
            # 设置标题字体
            elif para.style.name.startswith("Heading"):
                for run in para.runs:
                    run.font.name = WORD_SETTING["英文字体"]
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_SETTING["中文字体"])
